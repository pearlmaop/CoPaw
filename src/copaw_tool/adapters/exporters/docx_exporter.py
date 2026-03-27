from pathlib import Path
from typing import Optional
from copaw_tool.domain.model.strategy_ir import StrategyIR
from copaw_tool.domain.model.report_model import CompletenessReport


class DocxExporter:
    def export_pseudocode(self, pseudocode: str, output_path: Path, ir: Optional[StrategyIR] = None) -> Path:
        try:
            import docx
            document = docx.Document()
            document.add_heading("标准化伪代码", 0)
            if ir:
                document.add_paragraph(f"策略ID: {ir.strategy_id}")
                document.add_paragraph(f"规则数量: {len(ir.rules)}")
            document.add_paragraph(pseudocode)
            document.save(str(output_path))
            return output_path
        except ImportError:
            raise RuntimeError("python-docx not installed. Run: pip install python-docx")

    def export_report(self, report: CompletenessReport, output_path: Path) -> Path:
        try:
            import docx
            document = docx.Document()
            document.add_heading("逻辑完整性报告", 0)
            document.add_paragraph(f"策略ID: {report.strategy_id}")
            document.add_paragraph(f"规则数量: {report.total_rules}")
            document.add_paragraph(f"问题数量: {report.issue_count}")
            document.add_paragraph(f"覆盖度评分: {report.coverage_score:.1f}%")
            if report.issues:
                document.add_heading("检测到的问题", 1)
                for issue in report.issues:
                    document.add_paragraph(f"[{issue.severity.upper()}] {issue.message}")
            document.add_heading("结论", 1)
            document.add_paragraph(report.conclusion)
            document.save(str(output_path))
            return output_path
        except ImportError:
            raise RuntimeError("python-docx not installed. Run: pip install python-docx")
